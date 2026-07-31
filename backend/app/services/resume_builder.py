import os
from typing import Dict, List

from docx import Document
from docx.shared import Pt
from docx2pdf import convert


class ResumeBuilder:

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """
        Replace only the visible text while preserving
        all formatting, styles, spacing and template.
        """

        if not paragraph.runs:
           paragraph.add_run(new_text)
           return

        # Keep formatting from first run
        first_run = paragraph.runs[0]

        first_run.text = new_text

        # Clear remaining runs WITHOUT deleting formatting
        for run in paragraph.runs[1:]:
            run.text = ""
    def _all_paragraphs(self, doc):

        paragraphs = []

        for p in doc.paragraphs:
            paragraphs.append(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        paragraphs.append(p)

        return paragraphs

    def _normalize_skills(self, skills):

        result = []

        for skill in skills:

            if isinstance(skill, dict):

                result.append(
                    skill.get("name")
                    or skill.get("skill")
                    or skill.get("value")
                    or ""
                )

            else:
                result.append(str(skill))

        return result
    
    def _update_summary(self, paragraphs, summary):
    
        for i, p in enumerate(paragraphs):

            if "professional summary" in p.text.lower():

                if i + 1 < len(paragraphs):
                    self._replace_paragraph_text(
                        paragraphs[i + 1],
                        summary
                    )

                    return

    def _update_skills(self, paragraphs, skills):

        skills = self._normalize_skills(skills)

        for i, p in enumerate(paragraphs):

            if "technical skills" in p.text.lower():

                if i + 1 < len(paragraphs):

                    self._replace_paragraph_text(
                        paragraphs[i + 1],
                        " | ".join(self._normalize_skills(skills))
                    )

                return

    def _update_projects(self, paragraphs, projects):
    
        for project in projects:

            title = project.get("title", "")
            technologies = project.get("technologies", "")
            bullets = project.get("description", [])

            for i, p in enumerate(paragraphs):

                if title.lower() not in p.text.lower():
                   continue

                # Update project title
                self._replace_paragraph_text(
                    p,
                    title
                )

                # Technologies line
                if i + 1 < len(paragraphs):
                    self._replace_paragraph_text(
                        paragraphs[i + 1],
                         technologies
                    )

                bullet_lines = []

                for para in paragraphs[i + 2:]:

                    text = para.text.strip().lower()

                    if (
                        "education" in text
                        or "experience" in text
                        or "technical skills" in text
                        or "projects" in text
                        or "certification" in text
                        or "achievement" in text
                        or "professional summary" in text
                    ):
                        break

                    if para.text.strip():
                       bullet_lines.append(para)

                for idx in range(min(len(bullet_lines), len(bullets))):
                     self._replace_paragraph_text(
                         bullet_lines[idx],
                         "• " + bullets[idx]
                     )

                break
             
    def build_resume(
        self,
        input_docx_path: str,
        tailored_data: Dict,
        output_path: str
    ):

        doc = Document(input_docx_path)
        
        from docx.shared import Pt

        

        paragraphs = self._all_paragraphs(doc)
       
        summary = tailored_data.get("tailored_summary", "")

        skills = tailored_data.get("tailored_skills", [])

        projects = tailored_data.get("tailored_projects", [])

        # -----------------------------
        # Update Resume
        # -----------------------------

        self._update_summary(
            paragraphs,
            summary
        )

        self._update_skills(
            paragraphs,
            skills
        )

        self._update_projects(
            paragraphs,
            projects
        )

        # -----------------------------
        # Save DOCX
        # -----------------------------

        os.makedirs(
            os.path.dirname(output_path),
            exist_ok=True
        )

        doc.save(output_path)

        # -----------------------------
        # Generate PDF
        # -----------------------------

        try:

            pdf_path = output_path.replace(
                ".docx",
                ".pdf"
            )

            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            convert(
                output_path,
                pdf_path
            )

            print("✅ PDF Generated Successfully")

        except Exception as e:

            print("❌ PDF Conversion Failed:", e)

        return output_path
        